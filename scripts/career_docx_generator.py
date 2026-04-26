"""scripts/career_docx_generator.py — DOCX resume generator.

Reads state/content/resumes/cv-short.md and produces a formatted, editable Word document.
The DOCX is intended for finishing touches before saving as PDF.

Usage:
    .venv-docx/bin/python scripts/career_docx_generator.py [output_path]

Naming convention mirrors PDF generator: cv-{company}-{date}.docx
"""
from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
_CV_PATH = _REPO_ROOT / "state" / "content" / "resumes" / "cv-short.md"
_OUTPUT_DIR = _REPO_ROOT / "output" / "career"

_NAVY = RGBColor(0x1A, 0x35, 0x5E)
_DARK = RGBColor(0x22, 0x22, 0x22)
_GRAY = RGBColor(0x55, 0x55, 0x55)
_FONT = "Calibri"


def _add_border(paragraph) -> None:
    """Add a thin navy bottom border to a paragraph (section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A355E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _run(p, text: str, *, size: float = 10, bold: bool = False,
         italic: bool = False, color: RGBColor | None = None):
    r = p.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _inline(p, text: str, *, size: float = 10,
            color: RGBColor | None = None) -> None:
    """Add text with **bold** inline markers parsed into bold runs."""
    color = color or _DARK
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            _run(p, part[2:-2], size=size, bold=True, color=color)
        elif part:
            _run(p, part, size=size, color=color)


def _spacing(p, *, before: float = 0, after: float = 0) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def generate(cv_path: Path, output_path: Path) -> None:
    lines = cv_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # Default font for Normal style
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    after_h1 = False  # next non-empty line after H1 is the contact line

    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        # H1 → Name (large, navy, bold)
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            _run(p, line[2:].strip(), size=20, bold=True, color=_NAVY)
            _spacing(p, before=0, after=3)
            after_h1 = True
            continue

        # Contact line — always immediately follows H1
        if after_h1:
            p = doc.add_paragraph()
            _run(p, line.strip(), size=9, color=_GRAY)
            _spacing(p, before=0, after=6)
            after_h1 = False
            continue

        # H2 → Section heading (UPPERCASE + navy bottom border)
        if line.startswith("## "):
            p = doc.add_paragraph()
            _run(p, line[3:].strip().upper(), size=10, bold=True, color=_NAVY)
            _add_border(p)
            _spacing(p, before=8, after=3)
            continue

        # H3 → Job title | Company | Dates
        if line.startswith("### "):
            title = line[4:].strip()
            p = doc.add_paragraph()
            parts = title.split(" | ")
            for j, part in enumerate(parts):
                if j > 0:
                    _run(p, "  |  ", size=10, color=_GRAY)
                # j=0 role title: bold dark  |  j=1 employer: bold navy  |  j=2 dates: gray
                if j == 1:
                    _run(p, part, size=10, bold=True, color=_NAVY)
                elif j == 2:
                    _run(p, part, size=10, color=_GRAY)
                else:
                    _run(p, part, size=10, bold=True, color=_DARK)
            _spacing(p, before=6, after=2)
            continue

        # Bullet (manual bullet character for reliable rendering)
        if line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.17)
            _run(p, "\u2022  ", size=10, color=_DARK)
            _inline(p, line[2:].strip(), size=10)
            _spacing(p, before=1, after=1)
            continue

        # Plain paragraph (summary body, competencies line, etc.)
        p = doc.add_paragraph()
        _inline(p, line, size=10)
        _spacing(p, before=1, after=2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\u2713 {output_path}")


if __name__ == "__main__":
    today = date.today().strftime("%Y-%m-%d")
    # Usage: generator.py [output_path] [cv_source_path]
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else _OUTPUT_DIR / f"cv-google-{today}.docx"
    cv = Path(sys.argv[2]) if len(sys.argv) > 2 else _CV_PATH
    generate(cv, out)
