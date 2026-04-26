"""Generate visually styled DOCX for Google Senior TPM II, Applied AI resume.

Two-column layout: navy sidebar (contact/skills/education) + main content area.
Based on V4 visual template pattern from tmp/generate_resume_v4.py.

Usage:
    .venv-docx/bin/python scripts/career_docx_visual_google.py [output_path]
"""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

_REPO_ROOT = Path(__file__).resolve().parent.parent
_OUTPUT_DIR = _REPO_ROOT / "output" / "career"

# Palette
NAVY       = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT     = RGBColor(0x2E, 0x86, 0xAB)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT  = RGBColor(0x2D, 0x2D, 0x2D)
MED_GRAY   = RGBColor(0x88, 0x88, 0x88)
SIDEBAR_BG = "1B2A4A"
ACCENT_HEX = "2E86AB"
NAVY_HEX   = "1B2A4A"
FONT_MAIN  = "Calibri"
FONT_SIDE  = "Calibri"


# ---------------------------------------------------------------------------
# Cell helpers
# ---------------------------------------------------------------------------

def _cell_bg(cell, hex_color: str) -> None:
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_margins(cell, *, top=0, bottom=0, left=100, right=100) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _cell_no_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcB.append(el)
    tcPr.append(tcB)


def _table_no_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblB = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblB.append(el)
    tblPr.append(tblB)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def _para_spacing(p, *, before=0, after=0, line=None) -> None:
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line is not None:
        p.paragraph_format.line_spacing = Pt(line)


def _run(p, text: str, *, size: float = 9, bold=False, italic=False,
         color: RGBColor | None = None, font: str = FONT_MAIN) -> None:
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def _bottom_border(p, hex_color: str = ACCENT_HEX, sz: str = "4") -> None:
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Sidebar building blocks
# ---------------------------------------------------------------------------

def sb_heading(cell, text: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=10, after=2)
    _run(p, text.upper(), size=8.5, bold=True, color=ACCENT, font=FONT_SIDE)
    # thin rule
    p2 = cell.add_paragraph()
    _para_spacing(p2, before=0, after=3)
    _run(p2, "━" * 20, size=5.5, color=ACCENT, font=FONT_SIDE)


def sb_text(cell, text: str, *, size: float = 8.5, bold=False, italic=False) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=0, after=1, line=11)
    _run(p, text, size=size, bold=bold, italic=italic, color=WHITE, font=FONT_SIDE)


def sb_item(cell, text: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=1, after=1, line=11)
    _run(p, "▸  ", size=7, color=ACCENT, font=FONT_SIDE)
    _run(p, text, size=8, color=RGBColor(0xCC, 0xDA, 0xEB), font=FONT_SIDE)


def sb_tag_row(cell, tags: list[str]) -> None:
    """Render a row of small pill-like tags (space-separated)."""
    p = cell.add_paragraph()
    _para_spacing(p, before=1, after=2, line=12)
    for i, tag in enumerate(tags):
        if i:
            _run(p, "  ", size=7.5, color=RGBColor(0xCC, 0xDA, 0xEB), font=FONT_SIDE)
        _run(p, f"· {tag}", size=7.5, color=RGBColor(0xCC, 0xDA, 0xEB), font=FONT_SIDE)


# ---------------------------------------------------------------------------
# Main-column building blocks
# ---------------------------------------------------------------------------

def main_section(cell, text: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=10, after=3)
    _run(p, text.upper(), size=10, bold=True, color=NAVY, font=FONT_MAIN)
    _bottom_border(p)


def main_role(cell, title: str, company_dates: str, *, scope: str = "") -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=7, after=0)
    _run(p, title, size=9.5, bold=True, color=DARK_TEXT, font=FONT_MAIN)

    p2 = cell.add_paragraph()
    _para_spacing(p2, before=0, after=1)
    _run(p2, company_dates, size=8.5, italic=True, color=ACCENT, font=FONT_MAIN)

    if scope:
        p3 = cell.add_paragraph()
        _para_spacing(p3, before=0, after=2)
        _run(p3, scope, size=8, italic=True, color=MED_GRAY, font=FONT_MAIN)


def main_subrole(cell, text: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=4, after=1)
    _run(p, text, size=8.5, bold=True, color=DARK_TEXT, font=FONT_MAIN)


def main_bullet(cell, text: str, *, bold_prefix: str = "") -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=1, after=1, line=12)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    _run(p, "●  ", size=5, color=ACCENT, font=FONT_MAIN)
    if bold_prefix:
        _run(p, bold_prefix, size=8.5, bold=True, color=DARK_TEXT, font=FONT_MAIN)
    _run(p, text, size=8.5, color=DARK_TEXT, font=FONT_MAIN)


def main_summary(cell, text: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=4, after=4, line=12)
    _run(p, text, size=8.5, color=DARK_TEXT, font=FONT_MAIN)


def main_project_line(cell, title: str, body: str) -> None:
    p = cell.add_paragraph()
    _para_spacing(p, before=3, after=1, line=12)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    _run(p, "●  ", size=5, color=ACCENT, font=FONT_MAIN)
    _run(p, title + " — ", size=8.5, bold=True, color=DARK_TEXT, font=FONT_MAIN)
    _run(p, body, size=8.5, color=DARK_TEXT, font=FONT_MAIN)


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------

def build(output_path: Path) -> None:
    doc = Document()

    # Remove default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Page setup
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(0.3)
        sec.bottom_margin = Inches(0.3)
        sec.left_margin = Inches(0.3)
        sec.right_margin = Inches(0.3)

    # Two-column table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    _table_no_borders(table)

    row = table.rows[0]
    left  = row.cells[0]
    right = row.cells[1]
    left.width  = Inches(2.35)
    right.width = Inches(5.55)

    _cell_bg(left, SIDEBAR_BG)
    _cell_margins(left,  top=210, bottom=210, left=190, right=150)
    _cell_margins(right, top=210, bottom=210, left=210, right=150)
    _cell_no_borders(left)
    _cell_no_borders(right)
    left.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ==================================================================
    # LEFT SIDEBAR
    # ==================================================================

    # Name
    p = left.paragraphs[0]
    _para_spacing(p, before=0, after=1)
    _run(p, "VED\nMISHRA", size=22, bold=True, color=WHITE, font=FONT_SIDE)

    # Title tagline
    p2 = left.add_paragraph()
    _para_spacing(p2, before=2, after=8)
    _run(p2, "Senior Technical\nProgram Manager\n— AI Deployment", size=9, italic=True,
         color=ACCENT, font=FONT_SIDE)

    # Contact
    sb_heading(left, "Contact")
    sb_text(left, "📍  Seattle / Redmond, WA")
    sb_text(left, "✉  vedprakash.m@me.com")
    sb_text(left, "🔗  linkedin.com/in/ved01")
    sb_text(left, "🌐  vedprakash.net")

    # Education
    sb_heading(left, "Education")
    sb_text(left, "MBA", bold=True)
    sb_text(left, "UW Foster, 2023")
    p = left.add_paragraph(); _para_spacing(p, before=0, after=3)

    sb_text(left, "MCS — Data Science", bold=True)
    sb_text(left, "UIUC, 2020  ·  GPA 3.9")
    sb_text(left, "Tau Beta Pi", italic=True)
    p = left.add_paragraph(); _para_spacing(p, before=0, after=3)

    sb_text(left, "BE Computer Engineering", bold=True)
    sb_text(left, "Pune University, 2006")
    sb_text(left, "Rank #1, Nashik Division", italic=True)

    # Certifications
    sb_heading(left, "Certifications")
    sb_item(left, "PMP — PMI #2557777")
    sb_item(left, "PMI-ACP — PMI #2687477")
    sb_item(left, "AWS Solutions Architect")

    # Core AI Skills
    sb_heading(left, "AI / LLM Skills")
    sb_item(left, "LLM Eval Pipelines")
    sb_item(left, "Agent Health Dashboards")
    sb_item(left, "Golden Dataset Frameworks")
    sb_item(left, "Multi-Agent Orchestration")
    sb_item(left, "RAG Pipelines")
    sb_item(left, "Claude API · MCP")
    sb_item(left, "OpenAI APIs · LangChain")
    sb_item(left, "Pilot → Production")
    sb_item(left, "Hallucination Guardrails")

    # Program Management
    sb_heading(left, "Program Management")
    sb_item(left, "OP1/OP2 · OKRs · Agile")
    sb_item(left, "Portfolio Review Cadence")
    sb_item(left, "Executive Communication")
    sb_item(left, "Cross-Geo Delivery")
    sb_item(left, "Vendor Management")

    # Tools
    sb_heading(left, "Tools & Platforms")
    sb_item(left, "Azure (Storage · DevOps)")
    sb_item(left, "Python · SQL · R")
    sb_item(left, "Snowflake · Databricks")
    sb_item(left, "Azure DevOps · Pilotfish")
    sb_item(left, "Playwright · LangChain")

    # ==================================================================
    # RIGHT MAIN CONTENT
    # ==================================================================

    # Summary
    main_section(right, "Professional Summary")
    main_summary(right,
        "Senior Technical Program Manager with 20 years of experience and deep specialization in "
        "AI system delivery — from LLM evaluation pipelines and agent health monitoring to "
        "enterprise pilot-to-production deployments. Built Shiproom AI (RAG eval pipeline, "
        "live at Microsoft) and Artha (multi-agent personal OS with golden-dataset regression "
        "suite). Proven operational glue across complex AI deployments at scale: customer health "
        "metrics, eval frameworks, and reliability systems built from zero. Comfortable at "
        "startup velocity and Fortune 500 enterprise scale."
    )

    # Experience
    main_section(right, "Professional Experience")

    # Microsoft
    main_role(right,
        "Senior Technical Program Manager — Azure Storage",
        "Microsoft Corporation  ·  Redmond, WA  ·  Aug 2024 – Present",
        scope="$760M–$1.6B projected COGS impact · 34 regions · 100+ engineers · monthly CVP/GM reviews"
    )
    main_bullet(right,
        "Built Shiproom AI: RAG pipeline over Azure DevOps with real-time at-risk detection and "
        "hallucination-guarded summarization — tracks ~50 Big Rocks, generates CVP/GM briefings "
        "at <$1/run. Direct analog to Customer Health Hub: Resolution Rate, at-risk signals, "
        "program velocity at weekly cadence.",
        bold_prefix="Shiproom AI (RAG eval pipeline, production) — "
    )
    main_bullet(right,
        "One of 30 demos selected org-wide; presented to Azure Core CTO Marcus Fontoura.",
        bold_prefix="Azure Core AI Demo Day — "
    )
    main_bullet(right,
        "Delivery TPM for fleet-wide Pilotfish control plane migration (~4,100 clusters, ~11K "
        "machines, 34 regions); reduced deployment time ~4 hrs → ~70 min; 48% major release "
        "time reduction.",
        bold_prefix="XPF Ramp — "
    )
    main_bullet(right,
        "Decomposes xStore monolith for 2× release velocity; DD-XPF Phase 1 Canary go-live "
        "March 2026; daily cross-geo standups (Redmond + Shanghai).",
        bold_prefix="Rubik / DD-XPF — "
    )

    # Opendoor
    main_role(right,
        "Staff Technical Program Manager — Research & Data Science",
        "Opendoor Technologies  ·  Seattle, WA  ·  Aug 2022 – Aug 2024",
        scope="Sole TPM for entire Data & Insights org · CDO reporting · $5B–$15B iBuying platform"
    )
    main_bullet(right,
        "Co-authored enterprise Data Contract; SLA monitoring across 123 critical ML pipeline "
        "tables (4 tiers: Resolution Rate, Latency, Data Quality, Freshness) — direct "
        "operational analog to a Customer Health Hub. GitHub Danger CI blocking unapproved "
        "schema changes.",
        bold_prefix="Customer Health analog — "
    )
    main_bullet(right,
        "\"Lowest sev0/1 count of the year\" (CTO MBR, June 2024) — team previously tagged "
        "for 20%+ of company high-severity incidents.",
        bold_prefix="90% reduction in Sev0/1 incidents — "
    )
    main_bullet(right,
        "Hired and led 5-person DataOps team from zero; Snowflake→Databricks migration "
        "(8,846 tables, 6-person vendor team); SOX compliance for pricing engine.",
        bold_prefix="DataOps & migrations — "
    )

    # Amazon
    main_role(right,
        "Technical Program Manager — 3 Roles",
        "Amazon  ·  Seattle, WA  ·  Apr 2020 – Sep 2022"
    )
    main_subrole(right, "Amazon Style (Kaspian ML) — ML Personalization  ·  Mar–Aug 2022")
    main_bullet(right,
        "29 tech teams across S2 launch; tracked ML eval metrics weekly (~4.4% conversion, "
        "20% replenishment); 50+ post-launch defects managed. National press: GMA, "
        "Business Insider."
    )
    main_subrole(right, "Amazon Relay Load Board — ML Pricing Eval  ·  Apr 2020–Sep 2021")
    main_bullet(right,
        "First TPM on RLB (30K+ carriers, >1/3 Amazon Middle Mile). ML eval lifecycle: "
        "deviation reports per model change, anomaly guardrails, A/B testing. Pricing "
        "accuracy 79%→85%; carrier premium 16%→5%."
    )
    main_subrole(right, "Alexa Voice Services (Cobra/Verizon) — Agent Health  ·  Sep 2021–Mar 2022")
    main_bullet(right,
        "Resolved multi-agent skill isolation bugs (ASK · Panda · Multiagent PE); "
        "latency 2s → <500ms; 10,000-device production trial."
    )

    # SAP
    main_role(right,
        "Sr. Consultant → Cloud Architect → Principal TPM",
        "SAP  ·  Newtown Square, PA / Walldorf, Germany  ·  Feb 2010 – Apr 2020"
    )
    main_bullet(right,
        "Oracle-to-HANA migration: 20,000+ tenants, 15+ global DCs, 128 TB migrated; "
        "SAP Executive Board reporting. GBaaS global rollout across 4 DC regions (US, EU, "
        "CN, KSA). 240+ enterprise customers as CoE Cloud Architect (IBM, Disney, BofA, "
        "Walmart, AT&T)."
    )

    # Infosys
    main_role(right,
        "Senior Software Engineer",
        "Infosys Technologies  ·  Pune, India  ·  Nov 2006 – Feb 2010"
    )
    main_bullet(right, "4 zero-defect production releases for British Telecom and Telstra.")

    # Projects
    main_section(right, "Independent AI Projects")
    main_project_line(right,
        "Artha — Multi-Agent Personal OS (2024–Present)",
        "Production system, 60+ integrations, 20+ domains. Architecture: EAR (Episodic Agent "
        "Router), multi-LLM routing (Claude · Gemini), MCP integrations, vault-encrypted state. "
        "Golden-dataset regression suite (200+ tests) covering agent accuracy, hallucination "
        "detection, and latency — the eval discipline this role demands at enterprise scale."
    )
    main_project_line(right,
        "Shiproom AI — AI Executive Reporting (2025, Microsoft)",
        "RAG pipeline + Azure DevOps. Tracks ~50 Big Rocks; hallucination guardrails; "
        "briefings at <$1/run. Methodology: spec-based AI dev with accuracy validation loop."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ {output_path}")


if __name__ == "__main__":
    today = date.today().strftime("%Y-%m-%d")
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else \
        _OUTPUT_DIR / f"cv-google-stpm2-applied-ai-{today}-visual.docx"
    build(out)
