#!/usr/bin/env python3
# pii-guard: ignore-file — entity names / document titles from SharePoint; no personal data fields
"""scripts/connectors/msgraph_sharepoint.py — SharePoint / OneDrive document connector.

Implements the handler contract (fetch + health_check) used by sharepoint_kb_sync.py
and, optionally, pipeline.py.

Fetch strategy:
  - Configured sites: delta queries for incremental sync.
      First run → _fetch_initial (full drive scan + seed delta link).
      Subsequent runs → _fetch_delta (only changed items since last run).
  - /me/drive/sharedWithMe: cursor-based polling (last_seen_shared_at).
      Items returned newest-first; early-exit when cursor date is passed.

State is persisted atomically to state/connectors/sharepoint_docs_state.yaml.

Security ordering (R21/HIGH):
  pii_guard.scan() runs on every extracted text BEFORE yielding to callers.
  Callers must NEVER call kg.add_episode() with raw text that has not passed
  this gate. The gate is inside this module, so caller code is unconditionally safe.

Dedup (R13):
  _is_already_ingested() uses the state YAML as the early gate.
  SQLite documents table is written by sharepoint_kb_sync.py, not here.

Ref: specs/kb-graph.md §10.10–§10.12
"""
from __future__ import annotations

import hashlib
import logging
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

_SCRIPTS_DIR = Path(__file__).resolve().parent.parent
_ROOT_DIR    = _SCRIPTS_DIR.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

_log = logging.getLogger(__name__)

_SP_STATE_FILE = _ROOT_DIR / "state" / "connectors" / "sharepoint_docs_state.yaml"
_ALLOWED_EXTS  = frozenset({".md", ".txt", ".docx"})


# ---------------------------------------------------------------------------
# State persistence helpers
# ---------------------------------------------------------------------------

def _load_state() -> dict:
    """Load connector state (dedup + delta links + shared-with-me cursor)."""
    try:
        import yaml
        if _SP_STATE_FILE.exists():
            with open(_SP_STATE_FILE, encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
    except Exception as exc:
        _log.warning("Failed to load SharePoint state: %s", exc)
    return {}


def _save_connector_state_atomic(state: dict) -> None:
    """Persist state atomically (.yaml.tmp → rename pattern)."""
    import yaml
    _SP_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
    tmp = _SP_STATE_FILE.with_suffix(".yaml.tmp")
    try:
        with open(tmp, "w", encoding="utf-8") as f:
            yaml.safe_dump(state, f, default_flow_style=False, allow_unicode=True)
        tmp.replace(_SP_STATE_FILE)
    except Exception:
        try:
            tmp.unlink(missing_ok=True)
        except Exception:
            pass
        raise


def _is_already_ingested(drive_item_key: str, state: dict, etag: str | None = None) -> bool:
    """Check whether a drive item has been ingested. etag-first fast path.

    Returns True (skip) when drive_item_key is present AND the etag matches
    (or no etag provided). Returns False when etag changed → re-ingest.
    """
    bucket = state.get("ingested_items", {})
    if drive_item_key not in bucket:
        return False
    if etag is not None:
        old_etag = bucket[drive_item_key].get("etag")
        if old_etag and old_etag != etag:
            return False  # etag changed → must re-ingest
    return True


# ---------------------------------------------------------------------------
# Document text extraction
# ---------------------------------------------------------------------------

def _extract_text(filename: str, content_bytes: bytes) -> str:
    """Convert document bytes to plain text for KB extraction pipeline."""
    lower = filename.lower()
    if lower.endswith(".md") or lower.endswith(".txt"):
        return content_bytes.decode("utf-8", errors="replace")
    if lower.endswith(".docx"):
        return _extract_docx_bytes(content_bytes)
    return ""


def _extract_docx_bytes(content: bytes) -> str:
    """Convert .docx bytes to pseudo-markdown. Skips gracefully if python-docx absent."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        _log.warning(
            "python-docx not installed — .docx extraction skipped "
            "(install with: pip install artha[docx])"
        )
        return ""
    import io
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        _log.warning("Failed to parse .docx bytes: %s", exc)
        return ""

    lines: list[str] = []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 2
            lines.append(f"{'#' * level} {para.text}")
        elif para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Graph API helpers
# ---------------------------------------------------------------------------

def _download_content(token: str, drive_id: str, item_id: str) -> bytes:
    """Download raw file bytes from the Graph API /content endpoint."""
    try:
        import requests as _req
    except ImportError as exc:
        raise RuntimeError(
            "[msgraph_sharepoint] 'requests' package not installed — run: pip install requests"
        ) from exc

    from lib.msgraph import GRAPH_BASE, _get_headers
    url  = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/content"
    resp = _req.get(url, headers=_get_headers(token), timeout=60, allow_redirects=True)

    if resp.status_code == 302:
        # Follow redirect manually (some Graph content endpoints redirect to storage)
        resp = _req.get(resp.headers["Location"], timeout=60)
    resp.raise_for_status()
    return resp.content


def _fetch_initial(
    token: str,
    drive_id: str,
    paths: list[str] | None,
) -> tuple[list[dict], str | None]:
    """Full initial scan of a drive or specific folder paths.

    Returns (items, delta_link). delta_link is None when scanning specific paths
    (no delta support for path-scoped queries).
    """
    from lib.msgraph import _graph_get, _graph_get_paginated, _graph_get_full_url

    delta_link: str | None = None
    all_items: list[dict]  = []

    if paths:
        for path in paths:
            escaped = path.strip("/")
            items = _graph_get_paginated(token, f"/drives/{drive_id}/root:/{escaped}:/children")
            all_items.extend(items)
    else:
        # Delta query on drive root — collect all items and seed the delta link
        try:
            page = _graph_get(token, f"/drives/{drive_id}/root/delta")
            all_items.extend(page.get("value", []))
            next_url   = page.get("@odata.nextLink")
            delta_link = page.get("@odata.deltaLink")
            while next_url:
                page = _graph_get_full_url(token, next_url)
                all_items.extend(page.get("value", []))
                next_url   = page.get("@odata.nextLink")
                delta_link = page.get("@odata.deltaLink") or delta_link
        except Exception as exc:
            _log.warning("Delta query failed on drive %s — falling back to /children: %s", drive_id, exc)
            fallback = _graph_get_paginated(token, f"/drives/{drive_id}/root/children")
            all_items.extend(fallback)

    return all_items, delta_link


def _fetch_delta(token: str, delta_link: str) -> tuple[list[dict], str | None]:
    """Incremental delta fetch using a persisted delta link.

    Returns (changed_items, new_delta_link). Caller should persist new_delta_link.
    """
    from lib.msgraph import _graph_get_full_url

    all_items: list[dict]  = []
    new_delta_link: str | None = None
    current_url: str | None = delta_link

    while current_url:
        page = _graph_get_full_url(token, current_url)
        all_items.extend(page.get("value", []))
        current_url    = page.get("@odata.nextLink")
        new_delta_link = page.get("@odata.deltaLink") or new_delta_link

    return all_items, new_delta_link


# ---------------------------------------------------------------------------
# sharedWithMe polling (§10.11 Path B)
# ---------------------------------------------------------------------------

def _fetch_shared_with_me(token: str, state: dict) -> Iterator[dict]:
    """Yield docs shared with the user from /me/drive/sharedWithMe.

    Uses last_seen_shared_at as an early-exit cursor. Graph returns items
    newest-first; stops paginating once item dates fall at or before the cursor.
    First run: O(N_all_time). Subsequent runs: O(N_new_since_last_run).

    R21/HIGH: pii_guard.scan() runs on extracted text BEFORE yielding.
    State mutations on ingested_items and shared_with_me cursor happen
    inside this function; caller must persist state after iteration.
    """
    import pii_guard
    from lib.msgraph import _graph_get_paginated

    cursor       = state.get("shared_with_me", {}).get("last_seen_shared_at")
    newest_seen: str | None = None

    items = _graph_get_paginated(token, "/me/drive/sharedWithMe")
    for item in items:
        shared_at = item.get("shared", {}).get("sharedDateTime", "")

        # Early exit: items arrive newest-first; stop when we reach the previous cursor
        if cursor and shared_at and shared_at <= cursor:
            break

        if newest_seen is None or (shared_at and shared_at > newest_seen):
            newest_seen = shared_at

        remote = item.get("remoteItem")
        if not remote:
            continue

        drive_id = remote.get("parentReference", {}).get("driveId", "")
        item_id  = remote.get("id", "")
        if not drive_id or not item_id:
            continue

        etag           = item.get("eTag", "")
        drive_item_key = f"{drive_id}:{item_id}"

        if _is_already_ingested(drive_item_key, state, etag=etag):
            continue

        try:
            content_bytes = _download_content(token, drive_id, item_id)
        except Exception as exc:
            _log.warning("Failed to download shared item '%s': %s", item.get("name", "?"), exc)
            continue

        content_text = _extract_text(item.get("name", ""), content_bytes)
        if not content_text.strip():
            continue

        # R21/HIGH: PII scan BEFORE the caller can call add_episode()
        pii_found, pii_types = pii_guard.scan(content_text)
        if pii_found:
            _log.warning(
                "Skipping shared doc '%s': PII detected (%s)",
                item.get("name", "?"),
                ", ".join(pii_types.keys()),
            )
            continue

        content_hash = hashlib.sha256(content_bytes).hexdigest()
        shared_by    = (
            item.get("shared", {})
                .get("sharedBy", {})
                .get("user", {})
                .get("displayName", "")
        )

        yield {
            "id":            item_id,
            "name":          item.get("name", ""),
            "web_url":       item.get("webUrl", remote.get("webUrl", "")),
            "drive_item_id": drive_item_key,
            "content_text":  content_text,
            "content_hash":  content_hash,
            "etag":          etag,
            "shared_by":     shared_by,
            "modified_at":   item.get("lastModifiedDateTime", ""),
            "site_name":     "",
            "library_path":  remote.get("parentReference", {}).get("path", ""),
            "source":        "shared_with_me",
        }

        # Update dedup state for this item
        state.setdefault("ingested_items", {})[drive_item_key] = {
            "content_hash": content_hash,
            "ingested_at":  datetime.now(timezone.utc).isoformat(),
            "etag":         etag,
            "name":         item.get("name", ""),
        }

    # Advance cursor to the newest item we saw
    if newest_seen:
        state.setdefault("shared_with_me", {})["last_seen_shared_at"] = newest_seen


# ---------------------------------------------------------------------------
# Primary handler contract
# ---------------------------------------------------------------------------

def fetch(auth_context: dict, config: dict) -> Iterator[dict]:
    """Fetch documents from configured SharePoint sites and sharedWithMe.

    Implements the handler contract: yields clean dicts consumed by
    sharepoint_kb_sync.py. All docs have already passed pii_guard.scan().

    Delta queries provide incremental sync — only returns documents that
    changed since the last successful fetch per drive.
    """
    import pii_guard
    from lib.msgraph import _graph_get

    token = auth_context["access_token"]
    state = _load_state()
    delta_links: dict = state.setdefault("delta_links", {})

    # --- Configured sites ---------------------------------------------------
    for site_cfg in config.get("sites", []):
        site_path = site_cfg.get("site", "")
        if not site_path:
            continue

        try:
            site_info = _graph_get(token, f"/sites/{site_path}")
        except Exception as exc:
            _log.warning("Failed to fetch site '%s': %s", site_path, exc)
            continue

        site_id      = site_info.get("id", "")
        site_display = site_info.get("displayName", site_path)

        for lib_cfg in site_cfg.get("libraries", []):
            lib_name     = lib_cfg.get("name", "Documents")
            allowed_exts = set(lib_cfg.get("file_types", list(_ALLOWED_EXTS)))
            paths        = lib_cfg.get("paths") or []

            # Resolve the drive for this library
            try:
                drives_resp = _graph_get(token, f"/sites/{site_id}/drives")
                drive = next(
                    (d for d in drives_resp.get("value", []) if d.get("name") == lib_name),
                    None,
                )
                if drive is None:
                    _log.warning("Drive '%s' not found in site '%s'", lib_name, site_path)
                    continue
                drive_id = drive["id"]
            except Exception as exc:
                _log.warning("Failed to resolve drive '%s/%s': %s", site_path, lib_name, exc)
                continue

            delta_key  = f"{site_id}:{drive_id}"
            delta_link = delta_links.get(delta_key)

            try:
                if delta_link:
                    items, new_delta = _fetch_delta(token, delta_link)
                else:
                    items, new_delta = _fetch_initial(token, drive_id, paths or None)
            except Exception as exc:
                _log.warning("Failed to fetch drive '%s': %s", drive_id, exc)
                continue

            if new_delta:
                delta_links[delta_key] = new_delta

            for item in items:
                name = item.get("name", "")

                # Skip deleted items (graph delta includes deletion markers)
                if item.get("deleted"):
                    continue

                if not any(name.lower().endswith(ext) for ext in allowed_exts):
                    continue

                etag           = item.get("eTag", "")
                drive_item_key = f"{drive_id}:{item['id']}"

                if _is_already_ingested(drive_item_key, state, etag=etag):
                    continue

                try:
                    content_bytes = _download_content(token, drive_id, item["id"])
                except Exception as exc:
                    _log.warning("Failed to download '%s': %s", name, exc)
                    continue

                content_text = _extract_text(name, content_bytes)
                if not content_text.strip():
                    continue

                # R21/HIGH: PII scan BEFORE yielding
                pii_found, pii_types = pii_guard.scan(content_text)
                if pii_found:
                    _log.warning(
                        "Skipping '%s': PII detected (%s)",
                        name, ", ".join(pii_types.keys()),
                    )
                    continue

                content_hash = hashlib.sha256(content_bytes).hexdigest()

                doc: dict = {
                    "id":            item["id"],
                    "name":          name,
                    "site_name":     site_display,
                    "library_path":  item.get("parentReference", {}).get("path", ""),
                    "web_url":       item.get("webUrl", ""),
                    "content_text":  content_text,
                    "modified_by":   (
                        item.get("lastModifiedBy", {})
                            .get("user", {})
                            .get("displayName", "")
                    ),
                    "modified_at":   item.get("lastModifiedDateTime", ""),
                    "etag":          etag,
                    "content_hash":  content_hash,
                    "drive_item_id": drive_item_key,
                    "source":        "sharepoint",
                }

                yield doc

                # Update dedup state and persist after each page
                state.setdefault("ingested_items", {})[drive_item_key] = {
                    "content_hash": content_hash,
                    "ingested_at":  datetime.now(timezone.utc).isoformat(),
                    "etag":         etag,
                    "name":         name,
                }
                _save_connector_state_atomic(state)

    # --- sharedWithMe -------------------------------------------------------
    if config.get("include_shared_with_me", True):
        yield from _fetch_shared_with_me(token, state)
        _save_connector_state_atomic(state)

    # Update last_run timestamp
    state["last_run_at"] = datetime.now(timezone.utc).isoformat()
    _save_connector_state_atomic(state)


def health_check(auth_context: dict, config: dict) -> bool:
    """Verify Graph API connectivity and sufficient authentication."""
    try:
        from lib.msgraph import _graph_get
        token = auth_context["access_token"]
        me    = _graph_get(token, "/me", params={"$select": "displayName"})
        return bool(me.get("displayName"))
    except Exception as exc:
        _log.warning("SharePoint connector health_check failed: %s", exc)
        return False
